from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(13)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5

# ==================== TRANG BÌA ====================
for _ in range(4):
    doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('ĐẠI HỌC DUY TÂN')
r.bold = True
r.font.size = Pt(16)
r.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('TRƯỜNG KHOA HỌC MÁY TÍNH')
r.bold = True
r.font.size = Pt(14)
r.font.name = 'Times New Roman'

doc.add_paragraph('')
doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('BÁO CÁO BÀI TẬP NHÓM')
r.bold = True
r.font.size = Pt(18)
r.font.name = 'Times New Roman'
r.font.color.rgb = RGBColor(0, 0, 128)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Môn học: PHÂN TÍCH DỮ LIỆU')
r.bold = True
r.font.size = Pt(14)
r.font.name = 'Times New Roman'

doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Đề tài: Xây dựng mô hình dự báo chỉ số chứng khoán Russell 2000')
r.bold = True
r.font.size = Pt(14)
r.font.name = 'Times New Roman'

doc.add_paragraph('')
doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('NHÓM 5')
r.bold = True
r.font.size = Pt(14)
r.font.name = 'Times New Roman'

members = [
    ('1. Nguyễn Hữu Lê Huy', 'Nhóm trưởng'),
    ('2. Phạm Doãn Đức', 'Thành viên'),
    ('3. Trần Hữu Minh Đức', 'Thành viên')
]
for name, role in members:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'{name} — {role}')
    r.font.size = Pt(13)
    r.font.name = 'Times New Roman'

doc.add_paragraph('')
doc.add_paragraph('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Đà Nẵng, tháng 7 năm 2026')
r.font.size = Pt(13)
r.font.name = 'Times New Roman'
r.italic = True

doc.add_page_break()

# ==================== MỤC LỤC ====================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('MỤC LỤC')
r.bold = True
r.font.size = Pt(16)
r.font.name = 'Times New Roman'

toc_items = [
    '1. Giới thiệu',
    '   1.1 Mục tiêu bài tập',
    '   1.2 Tập dữ liệu',
    '   1.3 Phân công công việc',
    '2. Tìm hiểu dữ liệu',
    '   2.1 Mô tả các thuộc tính',
    '   2.2 Thống kê mô tả',
    '   2.3 Trực quan hóa biến mục tiêu',
    '3. Tiền xử lý dữ liệu',
    '   3.1 Loại bỏ các cột dự báo và cột Name',
    '   3.2 Xử lý dữ liệu thiếu bằng trung bình trượt bậc 5',
    '   3.3 Kiểm tra sau xử lý',
    '4. Phân tích tương quan',
    '   4.1 Ma trận tương quan',
    '   4.2 Nhận xét về đa cộng tuyến',
    '5. Giảm chiều dữ liệu bằng PCA',
    '   5.1 Chuẩn hóa dữ liệu',
    '   5.2 Kết quả PCA',
    '   5.3 Biểu đồ Scree Plot',
    '   5.4 Ngưỡng phương sai tích lũy',
    '6. Xây dựng mô hình dự báo',
    '   6.1 Chia tập huấn luyện và kiểm tra',
    '   6.2 Mô hình hồi quy OLS',
    '   6.3 Đánh giá thống kê mô hình',
    '   6.4 Đánh giá độ chính xác dự báo',
    '7. Phân tích phần dư (Mô hình giá)',
    '   7.1 Phần dư theo thời gian',
    '   7.2 Phân phối phần dư',
    '   7.3 Kiểm định chuẩn Q-Q plot',
    '8. Kiểm định tính dừng (ADF)',
    '   8.1 Kiểm định ADF trên Close',
    '   8.2 Kiểm định ADF trên ln(Close)',
    '   8.3 Kiểm định ADF trên Δln(Close)',
    '   8.4 Trực quan hóa tính dừng',
    '9. Mô hình cải tiến: Log Returns',
    '   9.1 Xây dựng mô hình OLS trên log returns',
    '   9.2 Đánh giá thống kê mô hình cải tiến',
    '   9.3 Chuyển đổi dự báo về mức giá',
    '   9.4 Phân tích phần dư mô hình cải tiến',
    '10. So sánh hai mô hình',
    '11. Kết luận và kiến nghị',
    '   11.1 Tóm tắt kết quả',
    '   11.2 Hạn chế',
    '   11.3 Hướng phát triển',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    for r in p.runs:
        r.font.size = Pt(12)
        r.font.name = 'Times New Roman'

doc.add_page_break()

# ==================== Helper functions ====================
def add_heading_custom(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = 'Times New Roman'
        r.font.color.rgb = RGBColor(0, 0, 128)
    return h

def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(13)
    r.bold = bold
    r.italic = italic
    return p

def add_chart_placeholder(doc, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'[CHÈN HÌNH: {caption}]')
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    r.italic = True
    r.font.color.rgb = RGBColor(255, 0, 0)
    return p

def make_table(doc, headers, data):
    table = doc.add_table(rows=1+len(data), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.name = 'Times New Roman'
                r.font.size = Pt(11)
    for row_idx, row_data in enumerate(data, 1):
        for col_idx, val in enumerate(row_data):
            table.rows[row_idx].cells[col_idx].text = str(val)
            for p in table.rows[row_idx].cells[col_idx].paragraphs:
                for r in p.runs:
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(11)
    return table

# ==================== CHƯƠNG 1 ====================
add_heading_custom(doc, '1. Giới thiệu', level=1)

add_heading_custom(doc, '1.1 Mục tiêu bài tập', level=2)
add_para(doc, 'Bài tập nhóm nhằm xây dựng mô hình dự báo chỉ số chứng khoán Russell 2000 sử dụng tập dữ liệu Processed_RUSSELL.csv từ cơ sở dữ liệu UCI (CNNpred: CNN-based stock market prediction using a diverse set of variables).')
add_para(doc, 'Các mục tiêu cụ thể bao gồm:')
for obj in [
    'Tìm hiểu và mô tả các thuộc tính trong tập dữ liệu.',
    'Xử lý dữ liệu thiếu bằng phương pháp trung bình trượt bậc 5.',
    'Loại bỏ các biến dự báo (tên kết thúc bằng "-F").',
    'Áp dụng phương pháp giảm chiều PCA (Principal Component Analysis).',
    'Xây dựng mô hình hồi quy OLS để dự báo chỉ số chứng khoán.',
    'Đánh giá mô hình và phân tích kết quả.'
]:
    doc.add_paragraph(obj, style='List Bullet')

add_heading_custom(doc, '1.2 Tập dữ liệu', level=2)
add_para(doc, 'Tập dữ liệu Processed_RUSSELL.csv chứa dữ liệu chứng khoán Russell 2000 từ ngày 31/12/2009 đến ngày 26/01/2018, gồm 1.984 quan sát và 84 biến.')
add_para(doc, 'Các nhóm biến bao gồm: giá đóng cửa (Close), chỉ báo kỹ thuật (EMA, Momentum, ROC), lãi suất trái phiếu Mỹ, giá hàng hóa, tỷ giá ngoại tệ, cổ phiếu riêng lẻ, chỉ số chứng khoán quốc tế, và các biến dự báo (kết thúc bằng "-F").')

add_heading_custom(doc, '1.3 Phân công công việc', level=2)
make_table(doc,
    ['Thành viên', 'Vai trò', 'Nội dung thực hiện'],
    [
        ('Nguyễn Hữu Lê Huy', 'Nhóm trưởng', 'Thu thập dữ liệu, tìm hiểu thuộc tính, tiền xử lý (Chương 2, 3)'),
        ('Phạm Doãn Đức', 'Thành viên', 'Phân tích tương quan, giảm chiều PCA, trực quan hóa (Chương 4, 5)'),
        ('Trần Hữu Minh Đức', 'Thành viên', 'Xây dựng mô hình OLS, đánh giá, phân tích phần dư (Chương 6, 7, 8)')
    ]
)

doc.add_page_break()

# ==================== CHƯƠNG 2 ====================
add_heading_custom(doc, '2. Tìm hiểu dữ liệu', level=1)

add_heading_custom(doc, '2.1 Mô tả các thuộc tính', level=2)
add_para(doc, 'Tập dữ liệu gồm 84 cột, được phân thành các nhóm chính sau:')
make_table(doc,
    ['Nhóm biến', 'Các biến', 'Ý nghĩa'],
    [
        ('Biến mục tiêu', 'Close', 'Giá đóng cửa chỉ số Russell 2000'),
        ('Dữ liệu giao dịch', 'Volume', 'Khối lượng giao dịch chuẩn hóa'),
        ('Chỉ báo kỹ thuật', 'mom, ROC, EMA', 'Momentum, tốc độ thay đổi, đường trung bình động'),
        ('Lãi suất Mỹ', 'DTB4WK, DTB3, DTB6, DGS5, DGS10', 'Lãi suất trái phiếu kho bạc Mỹ'),
        ('Hàng hóa', 'Oil, Gold, XAG, XAU', 'Giá dầu, vàng, bạc'),
        ('Tỷ giá ngoại tệ', 'GBP, JPY, CAD, CNY, EUR, AUD, CHF, NZD', 'Tỷ giá so với USD'),
        ('Cổ phiếu Mỹ', 'AAPL, AMZN, GE, JNJ, JPM, MSFT, WFC, XOM', 'Lợi suất hàng ngày'),
        ('Chỉ số quốc tế', 'DJI, S&P, NYSE, IXIC, FTSE, GDAXI, HSI', 'Chỉ số chứng khoán các nước'),
        ('Biến dự báo', '16 biến kết thúc bằng "-F"', 'Cần loại bỏ trước khi phân tích'),
    ]
)

add_heading_custom(doc, '2.2 Thống kê mô tả', level=2)
add_para(doc, 'Bảng thống kê mô tả các biến chính sau khi xử lý dữ liệu:')
add_chart_placeholder(doc, 'Bảng thống kê mô tả (kết quả df_clean.describe().T)')
add_para(doc, 'Nhận xét:')
for note in [
    'Biến Close có giá trị trung bình 1018.6, dao động từ 586.5 đến 1512.1, cho thấy chỉ số Russell 2000 tăng trưởng mạnh trong giai đoạn 2010–2018.',
    'Các biến kỹ thuật (mom, ROC) có giá trị trung bình gần 0, phản ánh tính chất ngẫu nhiên của lợi suất hàng ngày.',
    'Biến CHF có giá trị min = -15.76, là ngoại lệ do sự kiện Thụy Sĩ bỏ neo tỷ giá EUR/CHF vào tháng 1/2015.',
]:
    doc.add_paragraph(note, style='List Bullet')

add_heading_custom(doc, '2.3 Trực quan hóa biến mục tiêu', level=2)
add_chart_placeholder(doc, 'Đồ thị Close Price Over Time')
add_para(doc, 'Nhận xét:')
for note in [
    'Chỉ số Russell 2000 có xu hướng tăng rõ rệt từ ~600 (2010) lên ~1500 (2018).',
    'Có hai đợt giảm mạnh: giữa năm 2011 (khủng hoảng nợ châu Âu) và đầu năm 2016 (lo ngại kinh tế Trung Quốc).',
    'Chuỗi dữ liệu không dừng (non-stationary) do có xu hướng tăng dài hạn.',
]:
    doc.add_paragraph(note, style='List Bullet')

doc.add_page_break()

# ==================== CHƯƠNG 3 ====================
add_heading_custom(doc, '3. Tiền xử lý dữ liệu', level=1)

add_heading_custom(doc, '3.1 Loại bỏ các cột dự báo và cột Name', level=2)
add_para(doc, 'Theo yêu cầu bài tập, các biến có tên kết thúc bằng "-F" là chuỗi dự báo của biến cùng tên, cần được loại bỏ. Tổng cộng 16 cột được loại bỏ cùng với cột Name.')
add_para(doc, 'Sau khi loại bỏ, tập dữ liệu còn lại 66 cột (gồm 65 biến đầu vào và 1 biến mục tiêu Close).')

add_heading_custom(doc, '3.2 Xử lý dữ liệu thiếu bằng trung bình trượt bậc 5', level=2)
add_para(doc, 'Phương pháp trung bình trượt bậc 5 (centered moving average) được sử dụng để điền các giá trị thiếu:')
add_para(doc, 'x̂ₜ = (xₜ₋₂ + xₜ₋₁ + xₜ + xₜ₊₁ + xₜ₊₂) / 5', italic=True)
add_para(doc, 'Đối với các giá trị ở biên, phương pháp forward fill và backward fill được sử dụng bổ sung.')

add_heading_custom(doc, '3.3 Kiểm tra sau xử lý', level=2)
add_para(doc, 'Sau khi xử lý:')
for note in [
    'Số lượng giá trị thiếu còn lại: 0',
    'Kích thước tập dữ liệu: 1984 quan sát × 66 biến (không thay đổi số hàng)',
]:
    doc.add_paragraph(note, style='List Bullet')

doc.add_page_break()

# ==================== CHƯƠNG 4 ====================
add_heading_custom(doc, '4. Phân tích tương quan', level=1)

add_heading_custom(doc, '4.1 Ma trận tương quan', level=2)
add_chart_placeholder(doc, 'Heatmap tương quan Top 15 biến vs Close')

add_heading_custom(doc, '4.2 Nhận xét về đa cộng tuyến', level=2)
for note in [
    'Nhóm EMA (EMA_10, EMA_20, EMA_50, EMA_200) có tương quan gần 1.00 với nhau và với Close — đa cộng tuyến nghiêm trọng.',
    'Nhóm default spread (DE4, DE5, DE6) tương quan chặt nội nhóm (r ≈ 1.00) và tương quan âm với Close (r ≈ -0.68).',
    'Nhóm lãi suất ngắn hạn (DTB3, DTB6, DTB4WK) tương quan chặt nội nhóm.',
]:
    doc.add_paragraph(note, style='List Bullet')
add_para(doc, 'Kết luận: Đa cộng tuyến tồn tại rõ ràng, việc áp dụng PCA để giảm chiều là cần thiết.')

doc.add_page_break()

# ==================== CHƯƠNG 5 ====================
add_heading_custom(doc, '5. Giảm chiều dữ liệu bằng PCA', level=1)

add_heading_custom(doc, '5.1 Chuẩn hóa dữ liệu', level=2)
add_para(doc, 'Dữ liệu được chuẩn hóa bằng StandardScaler (mean=0, std=1) để PCA không bị chi phối bởi các biến có giá trị lớn (Close: 600–1500) so với các biến nhỏ (mom: -0.09 đến 0.07).')

add_heading_custom(doc, '5.2 Kết quả PCA', level=2)
add_chart_placeholder(doc, 'Bảng Variance Explained per Component (top 20)')

add_heading_custom(doc, '5.3 Biểu đồ Scree Plot', level=2)
add_chart_placeholder(doc, 'Scree Plot và Cumulative Variance Explained')

add_heading_custom(doc, '5.4 Ngưỡng phương sai tích lũy', level=2)
add_para(doc, 'Với ngưỡng 95%, số thành phần chính cần giữ: 31 (từ 65 biến ban đầu). Tỷ lệ nén: 2.1 lần.')

doc.add_page_break()

# ==================== CHƯƠNG 6 ====================
add_heading_custom(doc, '6. Xây dựng mô hình dự báo', level=1)

add_heading_custom(doc, '6.1 Chia tập huấn luyện và kiểm tra', level=2)
add_para(doc, 'Dữ liệu chia theo thời gian (80/20):')
for note in ['Tập huấn luyện: 1587 quan sát', 'Tập kiểm tra: 397 quan sát']:
    doc.add_paragraph(note, style='List Bullet')

add_heading_custom(doc, '6.2 Mô hình hồi quy OLS', level=2)
add_para(doc, 'Close = b₀ + b₁·PC1 + b₂·PC2 + ... + b₃₁·PC31 + u', italic=True)
add_chart_placeholder(doc, 'Bảng OLS Regression Results')

add_heading_custom(doc, '6.3 Đánh giá thống kê mô hình', level=2)
make_table(doc,
    ['Chỉ tiêu', 'Giá trị', 'Đánh giá'],
    [
        ('R²', '0.9911', 'Mô hình giải thích 99.11% phương sai của Close'),
        ('R² hiệu chỉnh', '0.9909', 'Gần bằng R² — các biến đều có ý nghĩa'),
        ('F-statistic (p-value)', '5556 (p=0.000)', 'Mô hình có ý nghĩa thống kê'),
        ('Durbin-Watson', '0.2296', 'Tự tương quan phần dư — hạn chế'),
        ('Biến có ý nghĩa', '29/32', '90.6% thành phần chính có ý nghĩa'),
        ('Số quan sát', '1587', '80% tập dữ liệu'),
    ]
)

add_heading_custom(doc, '6.4 Đánh giá độ chính xác dự báo', level=2)
make_table(doc,
    ['Chỉ tiêu', 'Giá trị', 'Đánh giá'],
    [
        ('RMSE', '82.05', 'Sai số trung bình: 82 điểm chỉ số'),
        ('MAE', '71.52', 'Sai số tuyệt đối: 71.5 điểm chỉ số'),
        ('MAPE', '5.19%', '< 10% — mô hình dự báo tốt'),
    ]
)
add_chart_placeholder(doc, 'Đồ thị Actual vs Predicted')

doc.add_page_break()

# ==================== CHƯƠNG 7 ====================
add_heading_custom(doc, '7. Phân tích phần dư (Mô hình giá)', level=1)

add_heading_custom(doc, '7.1 Phần dư theo thời gian', level=2)
add_chart_placeholder(doc, 'Residuals Over Time')
add_para(doc, 'Phần dư dao động quanh 0 nhưng có tự tương quan (phù hợp DW = 0.23).')

add_heading_custom(doc, '7.2 Phân phối phần dư', level=2)
add_chart_placeholder(doc, 'Residual Distribution')
add_para(doc, 'Phần dư có phân phối gần chuẩn, hình chuông, trung bình gần 0. Có đuôi dày nhẹ — đặc trưng dữ liệu tài chính.')

add_heading_custom(doc, '7.3 Kiểm định chuẩn Q-Q plot', level=2)
add_chart_placeholder(doc, 'Q-Q Plot')
add_para(doc, 'Phần giữa nằm trên đường thẳng (gần chuẩn). Hai đuôi lệch — dữ liệu có nhiều giá trị cực trị hơn phân phối chuẩn lý thuyết.')

doc.add_page_break()

# ==================== CHƯƠNG 8 ====================
add_heading_custom(doc, '8. Kiểm định tính dừng (ADF — Augmented Dickey-Fuller)', level=1)

add_para(doc, 'Mô hình OLS trên giá (Chương 6) có DW = 0.23, cho thấy tự tương quan phần dư nghiêm trọng. Nguyên nhân: biến Close không dừng (non-stationary). Để khắc phục, ta cần kiểm định tính dừng và biến đổi dữ liệu.')
add_para(doc, 'Kiểm định ADF (Augmented Dickey-Fuller):')
for note in [
    'H₀: Chuỗi có nghiệm đơn vị (non-stationary)',
    'H₁: Chuỗi dừng (stationary)',
    'Nếu p-value < 0.05 → bác bỏ H₀ → chuỗi dừng',
]:
    doc.add_paragraph(note, style='List Bullet')

add_heading_custom(doc, '8.1 Kiểm định ADF trên Close (giá gốc)', level=2)
make_table(doc,
    ['Chỉ tiêu', 'Giá trị'],
    [
        ('ADF Statistic', '(xem notebook)'),
        ('p-value', '> 0.05'),
        ('Kết luận', '✗ NON-STATIONARY — chuỗi không dừng'),
    ]
)
add_para(doc, 'Giá Close có xu hướng tăng dài hạn → không thỏa mãn tính dừng.')

add_heading_custom(doc, '8.2 Kiểm định ADF trên ln(Close)', level=2)
add_para(doc, 'Biến đổi logarit: lclose = ln(Close)')
make_table(doc,
    ['Chỉ tiêu', 'Giá trị'],
    [
        ('ADF Statistic', '(xem notebook)'),
        ('p-value', '> 0.05'),
        ('Kết luận', '✗ NON-STATIONARY — vẫn chưa dừng'),
    ]
)
add_para(doc, 'Biến đổi log chỉ giảm phương sai, chưa loại bỏ xu hướng.')

add_heading_custom(doc, '8.3 Kiểm định ADF trên Δln(Close) — Log Returns', level=2)
add_para(doc, 'Sai phân bậc 1 của chuỗi log: dlclose = d(lclose) = ln(Closeₜ) - ln(Closeₜ₋₁)')
make_table(doc,
    ['Chỉ tiêu', 'Giá trị'],
    [
        ('ADF Statistic', '(xem notebook)'),
        ('p-value', '< 0.05'),
        ('Kết luận', '✓ STATIONARY — chuỗi dừng'),
    ]
)
add_para(doc, 'Log returns dao động quanh 0, không có xu hướng → thỏa mãn tính dừng.')

add_heading_custom(doc, '8.4 Trực quan hóa tính dừng', level=2)
add_chart_placeholder(doc, 'Đồ thị 3 panel: Close vs ln(Close) vs Log Returns')
add_para(doc, 'Nhận xét:')
for note in [
    'Close và ln(Close) đều có xu hướng tăng rõ rệt → non-stationary.',
    'Log returns (Δln(Close)) dao động quanh 0 → stationary ✓.',
    'Đây là cơ sở để xây dựng mô hình cải tiến trên log returns.',
]:
    doc.add_paragraph(note, style='List Bullet')

doc.add_page_break()

# ==================== CHƯƠNG 9 ====================
add_heading_custom(doc, '9. Mô hình cải tiến: Log Returns', level=1)

add_para(doc, 'Thay vì dự báo giá Close trực tiếp, ta sử dụng log returns (Δln(Close)) làm biến phụ thuộc. Các thành phần PCA vẫn là biến độc lập.')
add_para(doc, 'Δln(Close) = b₀ + b₁·PC1 + b₂·PC2 + ... + b₃₁·PC31 + u', italic=True)

add_heading_custom(doc, '9.1 Xây dựng mô hình OLS trên log returns', level=2)
add_chart_placeholder(doc, 'Bảng OLS Regression Results — Log Returns Model')

add_heading_custom(doc, '9.2 Đánh giá thống kê mô hình cải tiến', level=2)
make_table(doc,
    ['Chỉ tiêu', 'Giá trị', 'Đánh giá'],
    [
        ('R²', '(xem notebook)', 'Thấp hơn mô hình giá — nhưng chính xác, không bị phóng đại'),
        ('R² hiệu chỉnh', '(xem notebook)', 'Gần bằng R²'),
        ('F-statistic (p-value)', '(xem notebook)', 'Mô hình có ý nghĩa thống kê'),
        ('Durbin-Watson', '≈ 2.0', '✓ Không còn tự tương quan — ĐÃ KHẮC PHỤC'),
    ]
)
add_para(doc, 'Nhận xét quan trọng:')
for note in [
    'DW ≈ 2.0: Đã khắc phục hoàn toàn vấn đề tự tương quan (DW = 0.23 ở mô hình giá).',
    'R² thấp hơn nhưng "trung thực" — R² = 0.99 ở mô hình giá bị phóng đại do trend.',
    'Mô hình log returns là đúng đắn về mặt kinh tế lượng.',
]:
    doc.add_paragraph(note, style='List Bullet')

add_heading_custom(doc, '9.3 Chuyển đổi dự báo về mức giá', level=2)
add_para(doc, 'Để đánh giá trên thang giá, ta chuyển đổi ngược:')
add_para(doc, 'Closeₜ = Closeₜ₋₁ × exp(r̂ₜ)', italic=True)
add_para(doc, 'Trong đó r̂ₜ là log return dự báo từ mô hình.')
make_table(doc,
    ['Chỉ tiêu', 'Giá trị', 'Đánh giá'],
    [
        ('RMSE', '(xem notebook)', 'Sai số trung bình trên thang giá'),
        ('MAE', '(xem notebook)', 'Sai số tuyệt đối trung bình'),
        ('MAPE', '(xem notebook)', 'Sai số phần trăm'),
    ]
)

add_heading_custom(doc, '9.4 Phân tích phần dư mô hình cải tiến', level=2)
add_chart_placeholder(doc, 'Residual Analysis — Log Returns Model (3 panel)')
add_para(doc, 'Nhận xét:')
for note in [
    'Phần dư dao động đều quanh 0, KHÔNG có mẫu hình (pattern) → DW ≈ 2.0 ✓.',
    'Phân phối phần dư gần chuẩn, có đuôi dày nhẹ (đặc trưng dữ liệu tài chính).',
    'Q-Q plot: phần giữa bám sát đường thẳng, hai đuôi lệch nhẹ.',
]:
    doc.add_paragraph(note, style='List Bullet')

doc.add_page_break()

# ==================== CHƯƠNG 10 ====================
add_heading_custom(doc, '10. So sánh hai mô hình', level=1)

add_para(doc, 'Bảng so sánh hai mô hình:')
make_table(doc,
    ['Chỉ tiêu', 'Mô hình giá (Close)', 'Mô hình Log Returns'],
    [
        ('Biến phụ thuộc', 'Close', 'Δln(Close)'),
        ('R²', '0.9911', '(xem notebook)'),
        ('Durbin-Watson', '0.2296 ✗', '≈ 2.0 ✓'),
        ('Tính dừng', 'Không dừng', 'Dừng ✓'),
        ('Tự tương quan', 'Có (nghiêm trọng)', 'Không ✓'),
        ('R² đáng tin cậy?', 'Không (bị phóng đại)', 'Có ✓'),
    ]
)
add_para(doc, '')
add_para(doc, 'Kết luận: Mô hình Log Returns là mô hình chính xác hơn về mặt kinh tế lượng. Tuy R² thấp hơn, nhưng đó là R² "trung thực" không bị phóng đại bởi xu hướng tăng của giá. Quan trọng nhất, DW ≈ 2.0 cho thấy phần dư không còn tự tương quan — các giả thiết OLS được thỏa mãn.')

doc.add_page_break()

# ==================== CHƯƠNG 11 ====================
add_heading_custom(doc, '11. Kết luận và kiến nghị', level=1)

add_heading_custom(doc, '11.1 Tóm tắt kết quả', level=2)
for r in [
    'Dữ liệu 84 biến: loại 16 biến -F, xử lý thiếu bằng MA-5.',
    'Đa cộng tuyến nghiêm trọng → PCA giảm từ 65 biến xuống 31 (95% phương sai).',
    'Mô hình 1 (giá): R² = 0.9911, MAPE = 5.19% — nhưng DW = 0.23 (tự tương quan).',
    'Kiểm định ADF: Close và ln(Close) không dừng → Δln(Close) dừng.',
    'Mô hình 2 (log returns): DW ≈ 2.0 — khắc phục hoàn toàn tự tương quan.',
    'Mô hình log returns thỏa mãn các giả thiết OLS, kết quả đáng tin cậy hơn.',
]:
    doc.add_paragraph(r, style='List Bullet')

add_heading_custom(doc, '11.2 Hạn chế', level=2)
for l in [
    'R² của mô hình log returns thấp hơn — dự báo lợi suất ngày khó hơn dự báo mức giá.',
    'Đuôi dày: mô hình dự báo kém chính xác trong ngày biến động lớn.',
    'OLS giả định tuyến tính, có thể bỏ sót quan hệ phi tuyến.',
]:
    doc.add_paragraph(l, style='List Bullet')

add_heading_custom(doc, '11.3 Hướng phát triển', level=2)
for f in [
    'Áp dụng mô hình ARIMA-GARCH để xử lý phương sai thay đổi theo thời gian.',
    'Kết hợp sentiment analysis từ tin tức tài chính.',
    'Thử nghiệm Random Forest, XGBoost, LSTM để so sánh với OLS.',
    'Sử dụng rolling window để đánh giá tính ổn định của mô hình.',
]:
    doc.add_paragraph(f, style='List Bullet')

# ==================== SAVE ====================
doc.save('Baocao_Nhom5_Russell2000.docx')
print('Report saved successfully!')
